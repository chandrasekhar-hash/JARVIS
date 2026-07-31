"""
JARVIS Product 1.6 - DOCX Parser.
Supports python-docx for Word document structure extraction.
"""

import os
import logging
from typing import Tuple, Dict, Any, Optional
from .base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    def __init__(self):
        super().__init__("DocxParser")

    def can_parse(self, source: str, mime_type: Optional[str] = None) -> bool:
        if mime_type and mime_type.lower() in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
            return True
        return source.lower().endswith(".docx") or source.lower().endswith(".doc")

    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Docx file not found: {file_path}")

        metadata = {"total_paragraphs": 0, "extracted_with": "none"}

        try:
            import docx  # type: ignore
            doc = docx.Document(file_path)
            paragraphs = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    paragraphs.append(p.text.strip())
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            metadata["total_paragraphs"] = len(paragraphs)
            metadata["extracted_with"] = "python-docx"
            return "\n\n".join(paragraphs), metadata
        except ImportError:
            logger.warning("python-docx not installed, using raw file reader fallback.")
        except Exception as e:
            logger.error(f"Docx parsing failed for {file_path}: {e}")

        with open(file_path, "r", errors="ignore") as f:
            content = f.read()
        return content, {"extracted_with": "raw_fallback"}
